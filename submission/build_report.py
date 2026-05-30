from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "Helping_Hands_Final_Report.docx"
SCHEMA = Path(__file__).resolve().parent / "Helping_Hands_Relational_Schema.png"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.add_run(text)
    return paragraph


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.10


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Helping Hands final report")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(102, 102, 102)


def build():
    doc = Document()
    configure_styles(doc)
    add_footer(doc.sections[0])

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Helping Hands")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string("111111")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Final project report - ELE 3921 Web Applications Development")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor.from_string("555555")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Team members: [add names]   |   GitHub: [add public repository link]   |   Video: [add video link]").italic = True

    add_heading(doc, "Project summary", 1)
    add_body(
        doc,
        "Helping Hands is a Django web application designed to connect seniors who need help with everyday tasks "
        "to trusted local helpers who want to earn money by helping their community. The problem addressed by the "
        "project is that many older adults can live independently, but still need occasional support with errands, "
        "technology, housework, transport, gardening, or small maintenance tasks. At the same time, students and "
        "younger people often want flexible local work. The website creates a simple marketplace around those needs."
    )
    add_body(
        doc,
        "The application has two main user roles: requesters and helpers. Requesters can create job posts, review "
        "applications, select a helper, message the selected person, mark jobs as completed, and leave reviews. "
        "Helpers can browse open jobs, apply with a short message, track pending applications, message requesters "
        "after selection, and build reputation through reviews."
    )

    add_heading(doc, "Main features", 1)
    add_bullets(
        doc,
        [
            "Registration and authentication using Django's built-in user system.",
            "Role-based accounts for requesters and helpers, with different dashboards and redirects.",
            "Job posting, category browsing, job details, applications, helper selection, and completion flow.",
            "Per-job messaging between the requester and selected helper.",
            "Profile pages with bio, location, avatar, posted/accepted jobs, and reviews.",
            "Review and report features after relevant job interactions.",
            "Admin interface for managing users, profiles, categories, jobs, applications, messages, reviews, and reports.",
            "Seeded demo data, uploaded avatars, Bootstrap, and a custom static stylesheet.",
        ],
    )

    add_heading(doc, "Architecture overview", 1)
    if SCHEMA.exists():
        doc.add_picture(str(SCHEMA), width=Inches(6.3))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run("Figure 1. Relational schema for the Helping Hands application.")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("555555")

    add_heading(doc, "Models", 2)
    add_table(
        doc,
        ["Model", "Purpose", "Important relationships"],
        [
            ["Category", "Stores job categories such as errands, cleaning, garden work, and tech help.", "Referenced by Job with protected deletion."],
            ["Profile", "Extends Django User with role, bio, location, and avatar.", "One-to-one with User."],
            ["Job", "Stores task posts, requester, helper, budget, timing, location, and status.", "Requester and helper connect to User; category connects to Category."],
            ["JobApplication", "Stores helper applications and optional application messages.", "Connects Job and applicant User; unique per job/applicant pair."],
            ["Message", "Stores messages in a job-specific conversation.", "Connects sender User to a Job."],
            ["Review", "Stores rating and comment after completed jobs.", "Connects reviewer, reviewee, and Job."],
            ["Report", "Stores reports for moderation and safety.", "Connects reporter, reported user, and related job."],
        ],
        [1.25, 2.85, 2.3],
    )

    add_heading(doc, "Views and templates", 2)
    add_body(
        doc,
        "The project uses Django templates for server-side rendering. Most user actions are implemented as regular "
        "Django views with redirects and messages after successful changes. The main views include:"
    )
    add_bullets(
        doc,
        [
            "Home page with different content for logged-in and logged-out users.",
            "Registration, login, and logout views.",
            "Profile, profile editing, and public user profile pages.",
            "Dashboard and completed jobs pages.",
            "Job list, job detail, job creation, application, helper selection, and job completion views.",
            "Message thread, review creation, and report creation views.",
        ],
    )

    add_heading(doc, "Forms and validation", 2)
    add_body(
        doc,
        "The application uses Django forms and model forms for user input. RegisterForm extends UserCreationForm and "
        "creates both the User and Profile. JobForm, ProfileForm, MessageForm, ReviewForm, and ReportForm handle the "
        "main user interactions. Form validation errors are displayed in the templates, while successful actions use "
        "Django messages to give feedback."
    )

    add_heading(doc, "Authentication and access control", 2)
    add_body(
        doc,
        "Authentication is based on Django's built-in User model. Views that require a logged-in user use login_required. "
        "Additional permission checks are handled in the views: only a requester can select a helper or complete their "
        "own job; only the requester and selected helper can access a message thread; reviews can only be left for "
        "completed jobs; and duplicate reviews/applications are prevented."
    )

    add_heading(doc, "Development process", 1)
    add_body(
        doc,
        "The project was developed around the core marketplace workflow. The models were created first so the main "
        "relationships were clear, then views and templates were added around the requester and helper journeys. "
        "The design aims to feel calm and trustworthy for seniors while still being easy for helpers to scan and use."
    )
    add_body(
        doc,
        "Most functionality is server-side Django, with only light JavaScript for interface details such as the role "
        "picker during registration and the message box scroll behavior. This kept the application simple and aligned "
        "with the course focus on Django, templates, forms, authentication, and database relationships."
    )

    add_heading(doc, "Front-end design", 1)
    add_body(
        doc,
        "The interface uses Bootstrap together with a custom static CSS file. The visual style uses a dark navigation "
        "bar, off-white page background, white content cards, terracotta accent buttons, serif headings, and readable "
        "form styling. The home page uses a large image hero, while the app pages are more functional and compact."
    )
    add_body(
        doc,
        "A custom static stylesheet is included at helply/static/helply/css/style.css. This was added so the project "
        "has committed static assets instead of relying only on inline CSS and external CDNs."
    )

    add_heading(doc, "Peer review reflection", 1)
    add_body(
        doc,
        "The peer review highlighted several strengths: the model design, the requester/helper split, access control, "
        "the admin setup, messaging, and the report feature. It also pointed out areas to improve before submission, "
        "especially the need for a clear requirements file, data dump, and custom static assets."
    )
    add_body(
        doc,
        "After reviewing the feedback, the project was cleaned up for submission. The repository now includes "
        "requirements.txt, a cleaner data dump, custom static CSS, and a README with setup instructions and sample "
        "credentials. JobApplication was also added to the admin interface so all important project data can be managed "
        "through Django admin."
    )

    add_heading(doc, "Future improvements", 1)
    add_bullets(
        doc,
        [
            "Add search and filtering by location, budget, time window, and review rating.",
            "Add notifications by email or in-app alerts when someone applies, is selected, or sends a message.",
            "Add stronger trust features such as identity verification, references, and admin moderation tools.",
            "Add payment handling or clear payment confirmation for real marketplace use.",
            "Improve mobile responsiveness further and add more automated tests for critical permission flows.",
        ],
    )

    add_heading(doc, "Sample user credentials", 1)
    add_body(doc, "The following demo accounts can be used to test the application. All use the password password123.")
    add_table(
        doc,
        ["Role", "Username", "Purpose"],
        [
            ["Admin", "louie", "Superuser account for Django admin."],
            ["Requester", "bjorn_h", "Senior/requester account for posting and managing jobs."],
            ["Requester", "ragnhild_s", "Requester account with example job activity."],
            ["Helper", "magnus_a", "Helper account for browsing and applying to jobs."],
            ["Helper", "soupy", "Additional helper account for testing role-specific flows."],
        ],
        [1.25, 1.5, 3.75],
    )

    add_heading(doc, "Submission links and files", 1)
    add_bullets(
        doc,
        [
            "GitHub repository: [add public repository link]",
            "Video presentation: [add video link]",
            "Data dump: project/data.json and project/fixtures/data.json",
            "Requirements file: project/requirements.txt",
            "Static CSS: project/helply/static/helply/css/style.css",
        ],
    )

    add_heading(doc, "AI tool usage", 1)
    add_body(
        doc,
        "AI assistance was used for review and cleanup support near the end of the project: comparing the assignment "
        "requirements with the project, identifying missing submission checklist items, moving existing CSS into a "
        "static file, registering a missing admin model, preparing cleaner submission notes, and drafting this report. "
        "The core project concept, Django architecture, models, views, and workflows should be understood and explained "
        "by the team during the presentation."
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
